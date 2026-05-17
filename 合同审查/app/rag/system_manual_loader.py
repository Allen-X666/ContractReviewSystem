"""
系统文档加载器

支持加载 Markdown 和 DOC/DOCX 格式的系统文档。
支持从文件路径或文件对象加载，直接返回文本内容。
"""

import os
import tempfile
import subprocess
import platform
from pathlib import Path
from typing import List, Optional, Union, BinaryIO

from langchain_community.document_loaders import UnstructuredMarkdownLoader


class SystemManualLoader:
    """
    系统文档加载器

    支持 Markdown 和 DOC/DOCX 格式的文档加载，
    直接返回文本字符串。

    注意：
    - .docx 格式使用 python-docx 库
    - .doc 格式（旧版 Word）使用以下方式（按优先级）：
      1. Windows: 使用 pywin32 调用 Word COM 接口
      2. Linux/Mac: 使用 antiword 命令行工具
      3. 使用 textract 库（如果已安装）
    """

    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding

    def load(
        self,
        file_path: Optional[Union[str, Path]] = None,
        file_obj: Optional[BinaryIO] = None,
        file_name: Optional[str] = None,
    ) -> str:
        """
        加载系统文档，返回完整文本内容

        Args:
            file_path: 文件路径（与 file_obj 二选一）
            file_obj: 文件对象（BytesIO 等，与 file_path 二选一）
            file_name: 文件名（当使用 file_obj 时必需，用于确定文件类型）

        Returns:
            str: 文档的完整文本内容
        """
        path, original_obj, name = self._get_file_info(file_path, file_obj, file_name)
        is_temp = file_obj is not None

        try:
            # 根据文件扩展名选择加载方式
            extension = Path(name).suffix.lower()

            if extension in ['.md', '.markdown']:
                # Markdown 文件使用 UnstructuredMarkdownLoader
                loader = UnstructuredMarkdownLoader(str(path), encoding=self.encoding)
                docs = loader.load()
                full_text = "\n\n".join([doc.page_content for doc in docs])
            elif extension == '.docx':
                # DOCX 文件使用 python-docx
                full_text = self._load_docx(str(path))
            elif extension == '.doc':
                # DOC 文件（旧版 Word）使用特殊处理
                full_text = self._load_doc(str(path))
            else:
                raise ValueError(f"不支持的文件格式: {extension}，仅支持 .md, .markdown, .docx, .doc")

            return full_text
        finally:
            if is_temp:
                self._cleanup_temp_file(path)

    def load_and_split(
        self,
        file_path: Optional[Union[str, Path]] = None,
        file_obj: Optional[BinaryIO] = None,
        file_name: Optional[str] = None,
    ) -> List[str]:
        """
        加载并拆分文档，返回文本块列表

        Args:
            file_path: 文件路径
            file_obj: 文件对象
            file_name: 文件名

        Returns:
            List[str]: 文本块列表
        """
        path, original_obj, name = self._get_file_info(file_path, file_obj, file_name)
        is_temp = file_obj is not None

        try:
            # 根据文件扩展名选择加载方式
            extension = Path(name).suffix.lower()

            if extension in ['.md', '.markdown']:
                # Markdown 文件使用 UnstructuredMarkdownLoader
                loader = UnstructuredMarkdownLoader(str(path), encoding=self.encoding)
                docs = loader.load()
                return [doc.page_content for doc in docs]
            elif extension == '.docx':
                # DOCX 文件使用 python-docx，按段落分块
                return self._load_docx_chunks(str(path))
            elif extension == '.doc':
                # DOC 文件加载完整文本后按段落拆分
                full_text = self._load_doc(str(path))
                # 按段落拆分（空行分隔）
                paragraphs = [p.strip() for p in full_text.split('\n') if p.strip()]
                return paragraphs
            else:
                raise ValueError(f"不支持的文件格式: {extension}，仅支持 .md, .markdown, .docx, .doc")
        finally:
            if is_temp:
                self._cleanup_temp_file(path)

    def _load_docx(self, file_path: str) -> str:
        """
        加载 DOCX 文件，返回完整文本

        Args:
            file_path: 文件路径

        Returns:
            str: 文档完整文本
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        # 使用 python-docx 加载文档
        doc = Document(file_path)

        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # 合并为一个完整文本
        full_text = "\n".join(paragraphs)

        return full_text

    def _load_docx_chunks(self, file_path: str) -> List[str]:
        """
        加载 DOCX 文件，按段落分块返回

        Args:
            file_path: 文件路径

        Returns:
            List[str]: 段落文本列表
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        # 使用 python-docx 加载文档
        doc = Document(file_path)

        # 提取所有段落文本
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        return paragraphs

    def _load_doc(self, file_path: str) -> str:
        """
        加载 DOC 文件（旧版 Word），返回完整文本

        尝试多种方式加载：
        1. Windows: 使用 pywin32 调用 Word COM 接口
        2. 使用 antiword 命令行工具
        3. 使用 textract 库
        4. 尝试转换为 docx 后读取

        Args:
            file_path: 文件路径

        Returns:
            str: 文档完整文本

        Raises:
            ValueError: 无法加载 DOC 文件
        """
        errors = []

        # 方法1: Windows 平台使用 pywin32
        if platform.system() == 'Windows':
            try:
                return self._load_doc_with_win32(file_path)
            except Exception as e:
                errors.append(f"pywin32: {e}")

        # 方法2: 使用 antiword 命令行工具
        try:
            return self._load_doc_with_antiword(file_path)
        except Exception as e:
            errors.append(f"antiword: {e}")

        # 方法3: 使用 textract 库
        try:
            return self._load_doc_with_textract(file_path)
        except Exception as e:
            errors.append(f"textract: {e}")

        # 方法4: 尝试使用 LibreOffice 转换为 docx
        try:
            return self._load_doc_with_libreoffice(file_path)
        except Exception as e:
            errors.append(f"libreoffice: {e}")

        # 所有方法都失败
        error_msg = "无法加载 .doc 文件，尝试的方法均失败:\n" + "\n".join(errors)
        error_msg += "\n\n建议解决方案:\n"
        error_msg += "1. Windows: 安装 pywin32: pip install pywin32\n"
        error_msg += "2. Linux: 安装 antiword: sudo apt-get install antiword\n"
        error_msg += "3. 安装 textract: pip install textract\n"
        error_msg += "4. 或将 .doc 文件手动转换为 .docx 格式后上传"
        raise ValueError(error_msg)

    def _load_doc_with_win32(self, file_path: str) -> str:
        """使用 Windows COM 接口加载 DOC 文件"""
        try:
            import win32com.client as win32
        except ImportError:
            raise ImportError("请安装 pywin32: pip install pywin32")

        word = None
        doc = None
        try:
            word = win32.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False

            # 打开文档
            doc = word.Documents.Open(str(file_path))

            # 提取文本
            paragraphs = []
            for para in doc.Paragraphs:
                text = para.Range.Text.strip()
                if text:
                    paragraphs.append(text)

            full_text = "\n".join(paragraphs)
            return full_text
        finally:
            if doc:
                doc.Close(SaveChanges=False)
            if word:
                word.Quit()

    def _load_doc_with_antiword(self, file_path: str) -> str:
        """使用 antiword 命令行工具加载 DOC 文件"""
        try:
            result = subprocess.run(
                ['antiword', file_path],
                capture_output=True,
                text=True,
                encoding=self.encoding,
                errors='replace'
            )
            if result.returncode != 0:
                raise RuntimeError(f"antiword 执行失败: {result.stderr}")
            return result.stdout
        except FileNotFoundError:
            raise RuntimeError("antiword 未安装")

    def _load_doc_with_textract(self, file_path: str) -> str:
        """使用 textract 库加载 DOC 文件"""
        try:
            import textract
        except ImportError:
            raise ImportError("请安装 textract: pip install textract")

        text = textract.process(file_path, encoding=self.encoding)
        return text.decode(self.encoding)

    def _load_doc_with_libreoffice(self, file_path: str) -> str:
        """使用 LibreOffice 将 DOC 转换为 DOCX 后读取"""
        import tempfile

        with tempfile.TemporaryDirectory() as tmpdir:
            # 构建输出路径
            output_path = os.path.join(tmpdir, "converted.docx")

            # 使用 LibreOffice 转换
            try:
                result = subprocess.run(
                    ['soffice', '--headless', '--convert-to', 'docx',
                     '--outdir', tmpdir, file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode != 0:
                    raise RuntimeError(f"LibreOffice 转换失败: {result.stderr}")

                # 检查转换后的文件是否存在
                if not os.path.exists(output_path):
                    raise RuntimeError("LibreOffice 转换后文件不存在")

                # 使用 python-docx 读取转换后的文件
                return self._load_docx(output_path)

            except FileNotFoundError:
                raise RuntimeError("LibreOffice 未安装")
            except subprocess.TimeoutExpired:
                raise RuntimeError("LibreOffice 转换超时")

    def _get_file_info(
        self,
        file_path: Optional[Union[str, Path]] = None,
        file_obj: Optional[BinaryIO] = None,
        file_name: Optional[str] = None,
    ) -> tuple[Path, Optional[BinaryIO], str]:
        """
        统一处理文件路径和文件对象

        Returns:
            tuple: (处理后的文件路径, 原始文件对象, 文件名)
        """
        if file_path is not None:
            path = Path(file_path)
            return path, None, path.name

        if file_obj is not None:
            if file_name is None:
                raise ValueError("使用 file_obj 时必须提供 file_name 参数")
            # 创建临时文件
            suffix = Path(file_name).suffix
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(file_obj.read())
                tmp_path = Path(tmp.name)
            # 重置文件对象指针
            file_obj.seek(0)
            return tmp_path, file_obj, file_name

        raise ValueError("必须提供 file_path 或 file_obj 之一")

    def _cleanup_temp_file(self, file_path: Path) -> None:
        """清理临时文件"""
        try:
            if file_path.exists():
                os.unlink(file_path)
        except Exception:
            pass


def load_system_manual(file_path: str) -> str:
    """
    便捷函数：加载系统文档

    Args:
        file_path: 文档文件路径（支持 .md, .markdown, .docx, .doc）

    Returns:
        str: 文档的完整文本内容
    """
    loader = SystemManualLoader()
    return loader.load(file_path=file_path)


def load_system_manual_chunks(file_path: str) -> List[str]:
    """
    便捷函数：加载系统文档并分块

    Args:
        file_path: 文档文件路径（支持 .md, .markdown, .docx, .doc）

    Returns:
        List[str]: 文本块列表
    """
    loader = SystemManualLoader()
    return loader.load_and_split(file_path=file_path)


if __name__ == "__main__":
    # 测试加载功能
    import sys

    # 测试 Markdown 文件
    md_file = r"e:\Professional\合同审查agent\data\system_operations\用户操作手册.md"

    # 测试 DOC 文件
    doc_file = r"e:\Professional\合同审查agent\合同审查\documents\新建 DOC 文档.doc"

    loader = SystemManualLoader()

    # 测试 Markdown
    print("=" * 50)
    print("测试 Markdown 文件:")
    print("=" * 50)
    try:
        text = loader.load(file_path=md_file)
        print(f"完整文本长度: {len(text)} 字符")
        print(f"前200字符预览:\n{text[:200]}...")
    except Exception as e:
        print(f"加载失败: {e}")

    # 测试 DOC
    print("\n" + "=" * 50)
    print("测试 DOC 文件:")
    print("=" * 50)
    try:
        text = loader.load(file_path=doc_file)
        print(f"完整文本长度: {len(text)} 字符")
        print(f"前200字符预览:\n{text[:200]}...")
    except Exception as e:
        print(f"加载失败: {e}")

    # 测试分块加载
    print("\n" + "=" * 50)
    print("测试 DOC 文件分块:")
    print("=" * 50)
    try:
        chunks = loader.load_and_split(file_path=doc_file)
        print(f"共分成 {len(chunks)} 个段落")
        for i, chunk in enumerate(chunks[:5]):
            print(f"\n段落 {i+1} (长度: {len(chunk)} 字符):")
            print(chunk[:100] + "..." if len(chunk) > 100 else chunk)
    except Exception as e:
        print(f"加载失败: {e}")
