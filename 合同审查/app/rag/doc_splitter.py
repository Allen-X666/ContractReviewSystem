"""
DOC/DOCX 文档拆分器

支持按段落、标题样式或分页符拆分 Word 文档。
"""

import re
from dataclasses import dataclass, field
from typing import List, Optional, Dict, Any, Union
from pathlib import Path


@dataclass
class DocSection:
    """Word 文档章节"""
    id: str
    title: str
    content: str
    style: Optional[str] = None  # 段落样式名称
    level: int = 0  # 标题级别（根据样式推断）
    metadata: Dict[str, Any] = field(default_factory=dict)


class DocSplitter:
    """
    Word 文档拆分器

    支持多种拆分策略：
    - 按段落拆分（每个段落作为一个条目）
    - 按标题样式拆分（根据 Heading 1/2/3 等样式）
    - 按分页符拆分
    """

    def __init__(
        self,
        split_by: str = "heading",  # 'heading', 'paragraph', 'page'
        max_depth: int = 3,  # 标题最大深度（仅 heading 模式）
        min_content_length: int = 50,
        merge_short_paragraphs: bool = True,  # 合并短段落
    ):
        self.split_by = split_by
        self.max_depth = max_depth
        self.min_content_length = min_content_length
        self.merge_short_paragraphs = merge_short_paragraphs

    def split(self, file_path: Union[str, Path]) -> List[DocSection]:
        """
        拆分 Word 文档

        Args:
            file_path: 文档路径

        Returns:
            List[DocSection]: 拆分后的章节列表
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        doc = Document(str(file_path))

        if self.split_by == "heading":
            return self._split_by_heading(doc)
        elif self.split_by == "paragraph":
            return self._split_by_paragraph(doc)
        elif self.split_by == "page":
            return self._split_by_page(doc)
        else:
            raise ValueError(f"不支持的拆分方式: {self.split_by}")

    def _split_by_heading(self, doc) -> List[DocSection]:
        """按标题样式拆分文档"""
        # 首先收集所有段落和标题信息
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            heading_match = re.match(r'Heading\s*(\d)', style_name, re.IGNORECASE)

            if heading_match:
                level = int(heading_match.group(1))
                paragraphs.append({
                    'type': 'heading',
                    'level': level,
                    'title': text,
                    'style': style_name,
                    'text': text,
                })
            else:
                paragraphs.append({
                    'type': 'paragraph',
                    'level': 0,
                    'title': '',
                    'style': style_name,
                    'text': text,
                })

        if not paragraphs:
            return []

        # 找出所有标题的位置
        headings = [(i, p) for i, p in enumerate(paragraphs) if p['type'] == 'heading']

        if not headings:
            # 没有标题，将整个文档作为一个章节
            full_text = '\n'.join([p['text'] for p in paragraphs])
            if full_text:
                return [DocSection(
                    id="1",
                    title="完整文档",
                    content=full_text,
                    style="Normal",
                    level=0,
                )]
            return []

        # 根据标题拆分章节
        sections = []
        for i, (idx, heading) in enumerate(headings):
            level = heading['level']

            # 只处理指定深度以内的标题
            if level > self.max_depth:
                continue

            # 确定当前章节的结束位置
            start_idx = idx
            if i + 1 < len(headings):
                # 找到下一个同等级或更高等级的标题
                end_idx = len(paragraphs)
                for next_idx, next_heading in headings[i + 1:]:
                    if next_heading['level'] <= level:
                        end_idx = next_idx
                        break
            else:
                end_idx = len(paragraphs)

            # 提取章节内容
            section_paras = paragraphs[start_idx:end_idx]
            content = '\n'.join([p['text'] for p in section_paras])

            # 检查内容长度
            if len(content) >= self.min_content_length:
                sections.append(DocSection(
                    id=str(len(sections) + 1),
                    title=heading['title'],
                    content=content,
                    style=heading['style'],
                    level=level,
                ))

        # 如果没有拆分到任何章节（可能都太短），将整个文档作为一个章节
        if not sections:
            full_text = '\n'.join([p['text'] for p in paragraphs])
            if full_text:
                sections.append(DocSection(
                    id="1",
                    title="完整文档",
                    content=full_text,
                    style="Normal",
                    level=0,
                ))

        return sections

    def _split_by_paragraph(self, doc) -> List[DocSection]:
        """按段落拆分文档"""
        sections = []
        current_content = []
        current_title = ""

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            # 检查是否是标题
            style_name = para.style.name if para.style else ""
            is_heading = style_name.startswith("Heading")

            if is_heading:
                # 保存之前的累积内容
                if current_content:
                    content = '\n'.join(current_content)
                    if len(content) >= self.min_content_length:
                        sections.append(DocSection(
                            id=str(len(sections) + 1),
                            title=current_title or "段落",
                            content=content,
                            style="Merged",
                            level=0,
                        ))
                    current_content = []
                current_title = text

            current_content.append(text)

            # 如果段落足够长，或者明确是标题后的第一个段落，作为一个独立条目
            if len(text) >= self.min_content_length or (is_heading and not self.merge_short_paragraphs):
                if len(current_content) > 1:  # 有累积的内容
                    content = '\n'.join(current_content[:-1])
                    if len(content) >= self.min_content_length:
                        sections.append(DocSection(
                            id=str(len(sections) + 1),
                            title=current_title or "段落",
                            content=content,
                            style="Normal",
                            level=0,
                        ))
                    current_content = [text]
                    current_title = text if is_heading else ""

        # 保存最后的内容
        if current_content:
            content = '\n'.join(current_content)
            if len(content) >= self.min_content_length:
                sections.append(DocSection(
                    id=str(len(sections) + 1),
                    title=current_title or "段落",
                    content=content,
                    style="Normal",
                    level=0,
                ))

        return sections if sections else self._split_by_heading(doc)

    def _split_by_page(self, doc) -> List[DocSection]:
        """按分页符拆分文档（简化实现，按段落累积估算）"""
        # python-docx 不直接支持分页符检测，这里使用段落数估算
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        if not paragraphs:
            return []

        # 每约 20 个段落作为一个页面
        page_size = 20
        sections = []

        for i in range(0, len(paragraphs), page_size):
            page_paras = paragraphs[i:i + page_size]
            content = '\n'.join(page_paras)

            if len(content) >= self.min_content_length:
                sections.append(DocSection(
                    id=str(len(sections) + 1),
                    title=f"第 {len(sections) + 1} 页",
                    content=content,
                    style="Page",
                    level=0,
                ))

        return sections


def split_doc(
    file_path: Union[str, Path],
    split_by: str = "heading",
    max_depth: int = 3,
) -> List[DocSection]:
    """
    便捷函数：拆分 Word 文档

    Args:
        file_path: 文档路径
        split_by: 拆分方式 ('heading', 'paragraph', 'page')
        max_depth: 标题最大深度

    Returns:
        List[DocSection]: 拆分后的章节列表
    """
    splitter = DocSplitter(
        split_by=split_by,
        max_depth=max_depth,
    )
    return splitter.split(file_path)


if __name__ == "__main__":
    # 测试 DOCX 拆分器
    test_file = r"e:\Professional\合同审查agent\合同审查\documents\新建 DOC 文档.docx"

    # 按标题拆分
    print("=" * 50)
    print("按标题样式拆分:")
    print("=" * 50)
    try:
        splitter = DocSplitter(split_by="heading", max_depth=3)
        sections = splitter.split(test_file)
        print(f"共拆分成 {len(sections)} 个章节:\n")
        for section in sections[:5]:
            print(f"ID: {section.id}")
            print(f"标题: {section.title}")
            print(f"样式: {section.style}")
            print(f"级别: {section.level}")
            print(f"内容长度: {len(section.content)} 字符")
            print(f"内容预览: {section.content[:100]}...")
            print("-" * 50)
    except Exception as e:
        print(f"拆分失败: {e}")

    # 按段落拆分
    print("\n" + "=" * 50)
    print("按段落拆分:")
    print("=" * 50)
    try:
        splitter = DocSplitter(split_by="paragraph")
        sections = splitter.split(test_file)
        print(f"共拆分成 {len(sections)} 个段落")
    except Exception as e:
        print(f"拆分失败: {e}")
